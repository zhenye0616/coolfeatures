"""Streamlit UI for the Document Template Pipeline.

Launch: uv run streamlit run app.py
"""

import json
import os
import shutil
import tempfile
from datetime import datetime
from pathlib import Path

import streamlit as st
from pdf2docx import Converter

from docx_template_pipeline import (
    LLMConfig,
    TemplateArtifacts,
    create_template,
    fill_template,
    generate_fill_prompt,
)

# ── Config ──────────────────────────────────────────────────────────────────

TEMPLATES_DIR = Path(__file__).parent / "templates"
TEMPLATES_DIR.mkdir(exist_ok=True)

st.set_page_config(
    page_title="DocForge",
    page_icon="◈",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ── Custom CSS ──────────────────────────────────────────────────────────────

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Libre+Baskerville:wght@400;700&family=Source+Sans+3:wght@300;400;500;600&display=swap');

/* ── Global ── */
.stApp {
    background: #FAFAF8;
}

section[data-testid="stSidebar"] {
    background: #1C1C1E;
}
section[data-testid="stSidebar"] * {
    color: #E5E5E3 !important;
}
section[data-testid="stSidebar"] label {
    color: #A0A09C !important;
    font-family: 'Source Sans 3', sans-serif !important;
    font-weight: 500 !important;
    letter-spacing: 0.03em;
    text-transform: uppercase;
    font-size: 0.7rem !important;
}
section[data-testid="stSidebar"] .stSelectbox > div > div,
section[data-testid="stSidebar"] input {
    background: #2C2C2E !important;
    border: 1px solid #3A3A3C !important;
    color: #E5E5E3 !important;
}

/* ── Header ── */
.docforge-header {
    padding: 2rem 0 1.5rem 0;
    border-bottom: 1px solid #D6D6D0;
    margin-bottom: 2rem;
}
.docforge-header h1 {
    font-family: 'Libre Baskerville', Georgia, serif;
    font-size: 1.8rem;
    font-weight: 700;
    color: #1C1C1E;
    margin: 0;
    letter-spacing: -0.01em;
}
.docforge-header p {
    font-family: 'Source Sans 3', sans-serif;
    font-size: 0.85rem;
    color: #7A7A74;
    margin: 0.35rem 0 0 0;
    font-weight: 400;
    letter-spacing: 0.02em;
}

/* ── Tabs ── */
.stTabs [data-baseweb="tab-list"] {
    gap: 0;
    border-bottom: 1px solid #D6D6D0;
    background: transparent;
}
.stTabs [data-baseweb="tab"] {
    font-family: 'Source Sans 3', sans-serif !important;
    font-weight: 600 !important;
    font-size: 0.85rem !important;
    letter-spacing: 0.06em !important;
    text-transform: uppercase !important;
    color: #7A7A74 !important;
    padding: 0.8rem 1.8rem !important;
    border-bottom: 2px solid transparent !important;
    background: transparent !important;
    transition: all 0.2s ease;
}
.stTabs [data-baseweb="tab"]:hover {
    color: #1C1C1E !important;
}
.stTabs [aria-selected="true"] {
    color: #1C1C1E !important;
    border-bottom: 2px solid #8B6914 !important;
    background: transparent !important;
}
.stTabs [data-baseweb="tab-highlight"] {
    display: none;
}
.stTabs [data-baseweb="tab-border"] {
    display: none;
}

/* ── Section labels ── */
.section-label {
    font-family: 'Source Sans 3', sans-serif;
    font-weight: 600;
    font-size: 0.7rem;
    letter-spacing: 0.1em;
    text-transform: uppercase;
    color: #8B6914;
    margin: 1.5rem 0 0.6rem 0;
}

/* ── Info cards ── */
.info-card {
    background: #F0EFEB;
    border: 1px solid #D6D6D0;
    border-radius: 6px;
    padding: 1rem 1.2rem;
    margin: 0.8rem 0;
    font-family: 'Source Sans 3', sans-serif;
    font-size: 0.85rem;
    color: #3C3C38;
    line-height: 1.5;
}
.info-card strong {
    color: #1C1C1E;
}

/* ── Template cards ── */
.template-card {
    background: #FFFFFF;
    border: 1px solid #D6D6D0;
    border-radius: 6px;
    padding: 1rem 1.2rem;
    margin: 0.4rem 0;
    transition: border-color 0.2s ease;
}
.template-card:hover {
    border-color: #8B6914;
}
.template-card h4 {
    font-family: 'Libre Baskerville', serif;
    font-size: 0.95rem;
    color: #1C1C1E;
    margin: 0 0 0.3rem 0;
}
.template-card .meta {
    font-family: 'Source Sans 3', sans-serif;
    font-size: 0.75rem;
    color: #7A7A74;
    letter-spacing: 0.02em;
}

/* ── Field badge ── */
.field-badge {
    display: inline-block;
    background: #F0EFEB;
    border: 1px solid #D6D6D0;
    border-radius: 3px;
    padding: 0.15rem 0.5rem;
    font-family: 'Source Sans 3', sans-serif;
    font-size: 0.75rem;
    font-weight: 500;
    color: #3C3C38;
    margin: 0.15rem 0.2rem 0.15rem 0;
}
.field-category {
    color: #8B6914;
    font-weight: 600;
}

/* ── Buttons ── */
.stButton > button[kind="primary"] {
    background: #1C1C1E !important;
    color: #FAFAF8 !important;
    border: none !important;
    font-family: 'Source Sans 3', sans-serif !important;
    font-weight: 600 !important;
    font-size: 0.8rem !important;
    letter-spacing: 0.06em !important;
    text-transform: uppercase !important;
    padding: 0.6rem 2rem !important;
    border-radius: 4px !important;
    transition: all 0.2s ease !important;
}
.stButton > button[kind="primary"]:hover {
    background: #8B6914 !important;
}
.stDownloadButton > button {
    background: #FFFFFF !important;
    color: #1C1C1E !important;
    border: 1px solid #1C1C1E !important;
    font-family: 'Source Sans 3', sans-serif !important;
    font-weight: 600 !important;
    font-size: 0.78rem !important;
    letter-spacing: 0.04em !important;
    text-transform: uppercase !important;
    padding: 0.5rem 1.5rem !important;
    border-radius: 4px !important;
    transition: all 0.2s ease !important;
}
.stDownloadButton > button:hover {
    background: #1C1C1E !important;
    color: #FAFAF8 !important;
}

/* ── Inputs ── */
.stTextArea textarea, .stTextInput input {
    font-family: 'Source Sans 3', sans-serif !important;
    font-size: 0.88rem !important;
    border: 1px solid #D6D6D0 !important;
    border-radius: 4px !important;
    background: #FFFFFF !important;
}
.stTextArea textarea:focus, .stTextInput input:focus {
    border-color: #8B6914 !important;
    box-shadow: 0 0 0 1px #8B6914 !important;
}

/* ── File uploader ── */
section[data-testid="stFileUploader"] {
    font-family: 'Source Sans 3', sans-serif !important;
}
section[data-testid="stFileUploader"] [data-testid="stFileUploaderDropzone"] {
    border: 1.5px dashed #C0BFB8 !important;
    border-radius: 6px !important;
    background: #FFFFFF !important;
    transition: border-color 0.2s ease !important;
}
section[data-testid="stFileUploader"] [data-testid="stFileUploaderDropzone"]:hover {
    border-color: #8B6914 !important;
}

/* ── Labels ── */
.stTextArea label, .stTextInput label, .stSelectbox label,
section[data-testid="stFileUploader"] label {
    font-family: 'Source Sans 3', sans-serif !important;
    font-weight: 500 !important;
    font-size: 0.78rem !important;
    letter-spacing: 0.05em !important;
    text-transform: uppercase !important;
    color: #5A5A56 !important;
}

/* ── Expander ── */
.streamlit-expanderHeader {
    font-family: 'Source Sans 3', sans-serif !important;
    font-weight: 600 !important;
    font-size: 0.8rem !important;
    letter-spacing: 0.04em !important;
    color: #3C3C38 !important;
}

/* ── Radio (for template selection) ── */
.stRadio label {
    text-transform: none !important;
}
.stRadio > div {
    gap: 0.3rem !important;
}

/* ── Status messages ── */
.stSuccess, .stInfo, .stWarning, .stError {
    font-family: 'Source Sans 3', sans-serif !important;
    font-size: 0.85rem !important;
    border-radius: 4px !important;
}

/* ── Divider ── */
hr {
    border-color: #E8E8E4 !important;
}

/* ── Streamlit footer hide ── */
footer {visibility: hidden;}
#MainMenu {visibility: hidden;}

/* ── Spinner ── */
.stSpinner > div {
    font-family: 'Source Sans 3', sans-serif !important;
    color: #5A5A56 !important;
}
</style>
""", unsafe_allow_html=True)

# ── Header ──────────────────────────────────────────────────────────────────

st.markdown("""
<div class="docforge-header">
    <h1>◈ DocForge</h1>
    <p>Preserve formatting. Replace data. Every line number, logo, and margin — intact.</p>
</div>
""", unsafe_allow_html=True)

# ── Sidebar: API Settings ───────────────────────────────────────────────────

with st.sidebar:
    st.markdown("#### API Configuration")

    provider = st.selectbox("Provider", ["OpenRouter", "Anthropic (direct)", "OpenAI-compatible"])

    default_key = os.environ.get("OPENROUTER_API_KEY", "") or os.environ.get("ANTHROPIC_API_KEY", "")
    api_key = st.text_input("API Key", type="password", value=default_key)

    provider_defaults = {
        "OpenRouter": ("https://openrouter.ai/api", "anthropic/claude-sonnet-4-20250514"),
        "Anthropic (direct)": ("", "claude-sonnet-4-20250514"),
        "OpenAI-compatible": ("", ""),
    }
    default_url, default_model = provider_defaults[provider]
    base_url = st.text_input("Base URL", value=default_url)
    model = st.text_input("Model", value=default_model)

    if not api_key:
        st.warning("Enter an API key to proceed.")


def get_config() -> LLMConfig:
    return LLMConfig(
        api_key=api_key,
        model=model,
        base_url=base_url if base_url.strip() else None,
    )


# ── Helpers ─────────────────────────────────────────────────────────────────


def convert_pdf_to_docx(pdf_path: str, docx_path: str) -> None:
    cv = Converter(pdf_path)
    cv.convert(docx_path)
    cv.close()


def list_saved_templates() -> list[dict]:
    """List all saved templates from the templates/ directory."""
    templates = []
    for d in sorted(TEMPLATES_DIR.iterdir()):
        if d.is_dir():
            analysis_path = d / "analysis.json"
            template_path = d / "template.docx"
            schema_path = d / "fill_schema.json"
            meta_path = d / "meta.json"
            if analysis_path.exists() and template_path.exists() and schema_path.exists():
                with open(analysis_path) as f:
                    analysis = json.load(f)
                with open(schema_path) as f:
                    fill_schema = json.load(f)
                # Read display name from meta.json if available
                display_name = d.name
                if meta_path.exists():
                    with open(meta_path) as f:
                        meta = json.load(f)
                    display_name = meta.get("name", d.name)
                templates.append({
                    "name": display_name,
                    "slug": d.name,
                    "path": d,
                    "analysis": analysis,
                    "fill_schema": fill_schema,
                    "template_path": str(template_path),
                    "doc_type": analysis.get("document_type", "Document"),
                    "field_count": len(analysis.get("fields", [])),
                })
    return templates


def save_template(name: str, artifacts: TemplateArtifacts) -> Path:
    """Save template artifacts to the templates/ directory."""
    slug = name.strip().replace(" ", "_").lower()
    dest = TEMPLATES_DIR / slug
    dest.mkdir(exist_ok=True)

    shutil.copy2(artifacts.template_path, dest / "template.docx")
    with open(dest / "analysis.json", "w") as f:
        json.dump(artifacts.analysis, f, indent=2)
    with open(dest / "fill_schema.json", "w") as f:
        json.dump(artifacts.fill_schema, f, indent=2)
    with open(dest / "meta.json", "w") as f:
        json.dump({
            "name": name,
            "document_type": artifacts.document_type,
            "document_description": artifacts.document_description,
            "field_count": artifacts.field_count,
            "placeholder_count": artifacts.placeholder_count,
            "created": datetime.now().isoformat(),
        }, f, indent=2)
    return dest


def load_artifacts(template_info: dict) -> TemplateArtifacts:
    """Load TemplateArtifacts from a saved template."""
    from docx import Document as DocxDocument
    import re

    template_path = template_info["template_path"]
    analysis = template_info["analysis"]
    fill_schema = template_info["fill_schema"]

    doc = DocxDocument(template_path)
    all_parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_parts.append(cell.text)
    for section in doc.sections:
        for p in section.header.paragraphs + section.footer.paragraphs:
            all_parts.append(p.text)
    placeholder_re = re.compile(r"\{\{\s*(\w+)\s*\}\}")
    placeholder_count = len(set(placeholder_re.findall("\n".join(all_parts))))

    return TemplateArtifacts(
        analysis=analysis,
        fill_schema=fill_schema,
        template_path=template_path,
        document_type=analysis.get("document_type", "Document"),
        document_description=analysis.get("document_description", ""),
        field_count=len(analysis.get("fields", [])),
        placeholder_count=placeholder_count,
    )


# ── Main Layout ─────────────────────────────────────────────────────────────

tab_generate, tab_create = st.tabs(["GENERATE DOCUMENT", "CREATE TEMPLATE"])

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  TAB 1: GENERATE DOCUMENT
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

with tab_generate:
    saved = list_saved_templates()

    if not saved and not api_key:
        st.markdown("""
        <div class="info-card">
            <strong>No templates yet.</strong> Go to the <strong>Create Template</strong> tab to
            extract a reusable template from any .docx or .pdf, then come back here to generate
            new documents from it.
        </div>
        """, unsafe_allow_html=True)

    else:
        col_left, col_right = st.columns([1, 1], gap="large")

        with col_left:
            st.markdown('<div class="section-label">1 — Choose template</div>', unsafe_allow_html=True)

            source_mode = st.radio(
                "Template source",
                options=(["Saved template"] if saved else []) + ["Upload .docx (one-shot)"],
                label_visibility="collapsed",
                horizontal=True,
            )

            selected_artifacts = None

            if source_mode == "Saved template" and saved:
                template_names = [t["name"] for t in saved]
                chosen_name = st.selectbox(
                    "Template",
                    template_names,
                    format_func=lambda n: f"{n}  ({next((t['field_count'] for t in saved if t['name'] == n), '?')} fields)",
                )
                chosen = next(t for t in saved if t["name"] == chosen_name)

                # Show fields summary
                fields_html = ""
                categories = {}
                for field in chosen["analysis"].get("fields", []):
                    cat = field.get("category", "other")
                    categories.setdefault(cat, []).append(field["field_name"])
                for cat, fnames in categories.items():
                    badges = "".join(f'<span class="field-badge">{fn}</span>' for fn in fnames)
                    fields_html += f'<div style="margin:0.4rem 0"><span class="field-badge field-category">{cat}</span> {badges}</div>'

                with st.expander(f"{chosen['field_count']} variable fields detected"):
                    st.markdown(fields_html, unsafe_allow_html=True)

                selected_artifacts = load_artifacts(chosen)

            else:
                uploaded_source = st.file_uploader(
                    "Upload source document",
                    type=["docx", "pdf"],
                    key="gen_upload",
                    help="Upload any .docx or .pdf — the template will be extracted automatically.",
                )

        with col_right:
            st.markdown('<div class="section-label">2 — Enter new case data</div>', unsafe_allow_html=True)

            new_facts = st.text_area(
                "Case facts",
                height=280,
                placeholder=(
                    "Attorney: Maria Santos (SBN 445566)\n"
                    "Firm: Santos Legal Group, APC\n"
                    "Plaintiff: Pacific Supply Co.\n"
                    "Defendant: Golden Gate Logistics LLC\n"
                    "Case No: 24STCV12345\n"
                    "Hearing: March 15, 2026 at 8:30 AM\n"
                    "..."
                ),
                help="Plain English. The LLM will map these facts to the template fields.",
            )

        # ── Generate button ──
        st.markdown("")  # spacer

        can_generate = bool(api_key and new_facts.strip())
        if source_mode == "Saved template":
            can_generate = can_generate and selected_artifacts is not None
        else:
            can_generate = can_generate and ("gen_upload" in st.session_state and st.session_state.gen_upload is not None)

        if st.button("Generate Document", type="primary", disabled=not can_generate, use_container_width=True):
            with tempfile.TemporaryDirectory() as tmpdir:
                # If one-shot mode: create template first
                if source_mode != "Saved template":
                    uploaded_file = st.session_state.gen_upload
                    source_path = os.path.join(tmpdir, uploaded_file.name)
                    with open(source_path, "wb") as f:
                        f.write(uploaded_file.getvalue())

                    # PDF conversion if needed
                    if uploaded_file.name.lower().endswith(".pdf"):
                        docx_path = source_path.rsplit(".", 1)[0] + ".docx"
                        with st.spinner("Converting PDF to DOCX..."):
                            convert_pdf_to_docx(source_path, docx_path)
                        source_path = docx_path

                    with st.spinner("Analyzing document structure..."):
                        selected_artifacts = create_template(
                            source_docx=source_path,
                            output_dir=tmpdir,
                            llm_config=get_config(),
                        )

                # Fill template
                tmp_template = os.path.join(tmpdir, "template_fill.docx")
                with open(selected_artifacts.template_path, "rb") as f:
                    tpl_bytes = f.read()
                with open(tmp_template, "wb") as f:
                    f.write(tpl_bytes)

                fill_artifacts = TemplateArtifacts(
                    analysis=selected_artifacts.analysis,
                    fill_schema=selected_artifacts.fill_schema,
                    template_path=tmp_template,
                    document_type=selected_artifacts.document_type,
                    document_description=selected_artifacts.document_description,
                    field_count=selected_artifacts.field_count,
                    placeholder_count=selected_artifacts.placeholder_count,
                )

                output_path = os.path.join(tmpdir, "generated.docx")
                with st.spinner("Generating document..."):
                    result = fill_template(
                        artifacts=fill_artifacts,
                        new_facts=new_facts,
                        output_path=output_path,
                        llm_config=get_config(),
                    )

                with open(output_path, "rb") as f:
                    doc_bytes = f.read()

            st.session_state["generated_doc"] = doc_bytes
            st.session_state["generated_result"] = result

        # ── Show result ──
        if "generated_doc" in st.session_state:
            st.markdown("")
            result = st.session_state["generated_result"]
            st.success(f"Done — {result.placeholder_count} fields replaced.")

            col_dl, col_detail = st.columns([1, 2])
            with col_dl:
                st.download_button(
                    "Download Document",
                    data=st.session_state["generated_doc"],
                    file_name="generated_document.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            with col_detail:
                with st.expander("View filled data"):
                    for k, v in sorted(result.context.items()):
                        st.markdown(f"`{k}` → **{v}**")


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  TAB 2: CREATE TEMPLATE
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

with tab_create:
    col_upload, col_info = st.columns([1, 1], gap="large")

    with col_upload:
        st.markdown('<div class="section-label">Upload source document</div>', unsafe_allow_html=True)
        uploaded = st.file_uploader(
            "Upload a .docx or .pdf",
            type=["docx", "pdf"],
            key="tpl_upload",
            help="The original document to extract a template from. All formatting is preserved.",
        )

        template_name = st.text_input(
            "Template name",
            placeholder="e.g. Motion to Quash Subpoena",
            help="A short name to identify this template later.",
        )

    with col_info:
        st.markdown('<div class="section-label">How it works</div>', unsafe_allow_html=True)
        st.markdown("""
        <div class="info-card">
            <strong>1.</strong> Upload any formatted .docx or .pdf<br>
            <strong>2.</strong> The LLM identifies all variable fields (names, dates, addresses...)<br>
            <strong>3.</strong> A reusable template is saved — every line number, logo, margin, and style preserved<br><br>
            Use the template in <strong>Generate Document</strong> to produce new documents with different data
            but identical formatting.
        </div>
        """, unsafe_allow_html=True)

    st.markdown("")

    can_create = bool(api_key and uploaded and template_name and template_name.strip())

    if st.button("Extract Template", type="primary", disabled=not can_create, use_container_width=True):
        with tempfile.TemporaryDirectory() as tmpdir:
            source_path = os.path.join(tmpdir, uploaded.name)
            with open(source_path, "wb") as f:
                f.write(uploaded.getvalue())

            if uploaded.name.lower().endswith(".pdf"):
                docx_path = source_path.rsplit(".", 1)[0] + ".docx"
                with st.spinner("Converting PDF to DOCX..."):
                    try:
                        convert_pdf_to_docx(source_path, docx_path)
                    except Exception as e:
                        st.error(f"PDF conversion failed: {e}")
                        st.stop()
                source_path = docx_path

            with st.spinner("Analyzing document with LLM — identifying all variable fields..."):
                artifacts = create_template(
                    source_docx=source_path,
                    output_dir=tmpdir,
                    llm_config=get_config(),
                )

            save_path = save_template(template_name.strip(), artifacts)

            st.session_state["last_created"] = {
                "name": template_name.strip(),
                "artifacts": artifacts,
                "path": str(save_path),
            }

    # ── Show result ──
    if "last_created" in st.session_state:
        info = st.session_state["last_created"]
        artifacts = info["artifacts"]

        st.success(f"Template **{info['name']}** saved — {artifacts.field_count} fields, {artifacts.placeholder_count} placeholders.")

        # Fields display
        fields_html = ""
        categories = {}
        for field in artifacts.analysis.get("fields", []):
            cat = field.get("category", "other")
            categories.setdefault(cat, []).append(field)

        for cat, fields in categories.items():
            badges = ""
            for f in fields:
                badges += f'<span class="field-badge" title="{f["description"]}">{f["field_name"]}</span>'
            fields_html += f'<div style="margin:0.5rem 0"><span class="field-badge field-category">{cat}</span> {badges}</div>'

        with st.expander("Detected fields", expanded=True):
            st.markdown(fields_html, unsafe_allow_html=True)

        col_dl1, col_dl2 = st.columns(2)
        with col_dl1:
            with open(Path(info["path"]) / "template.docx", "rb") as f:
                st.download_button(
                    "Download Template (.docx)",
                    data=f.read(),
                    file_name="template.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
        with col_dl2:
            st.download_button(
                "Download Schema (.json)",
                data=json.dumps(artifacts.fill_schema, indent=2),
                file_name="fill_schema.json",
                mime="application/json",
                use_container_width=True,
            )

    # ── Saved templates library ──
    saved = list_saved_templates()
    if saved:
        st.markdown("")
        st.markdown('<div class="section-label">Saved templates</div>', unsafe_allow_html=True)

        for t in saved:
            meta_path = t["path"] / "meta.json"
            created = ""
            if meta_path.exists():
                with open(meta_path) as f:
                    meta = json.load(f)
                created = meta.get("created", "")[:10]

            st.markdown(f"""
            <div class="template-card">
                <h4>{t["name"]}</h4>
                <div class="meta">{t["doc_type"]} &nbsp;·&nbsp; {t["field_count"]} fields{f" &nbsp;·&nbsp; {created}" if created else ""}</div>
            </div>
            """, unsafe_allow_html=True)
