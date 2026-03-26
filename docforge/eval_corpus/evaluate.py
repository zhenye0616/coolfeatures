"""
DocForge evaluation script.

Produces a single scalar score (0.0-1.0) measuring template extraction quality.
Analogous to autoresearch's evaluate_bpb in prepare.py — this is the fixed metric.

Usage:
    uv run evaluate.py                          # full eval, all corpus entries
    uv run evaluate.py --fast                   # fast check only (no LLM, no score)
    uv run evaluate.py --subset 3               # eval first N corpus entries
"""

import argparse
import json
import os
import sys
import tempfile
import time
from dataclasses import dataclass

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document

# Import the pipeline under test
from docx_template_pipeline import (
    LLMConfig,
    TemplateArtifacts,
    create_template,
    fill_template,
    _PLACEHOLDER_RE,
)


# ── Constants (DO NOT MODIFY — this is the fixed evaluation) ──────────────

CORPUS_DIR = os.path.join(os.path.dirname(__file__), "corpus")

WEIGHT_RECALL = 0.35
WEIGHT_PRECISION = 0.20
WEIGHT_INTEGRITY = 0.15
WEIGHT_ROUNDTRIP = 0.30

BASELINE_LINES = 712  # docx_template_pipeline.py line count at baseline
BLOAT_THRESHOLD = 1.5  # penalty kicks in at 1.5x baseline
BLOAT_PENALTY = 0.95   # 5% score reduction for bloated code


# ── Data structures ──────────────────────────────────────────────────────

@dataclass
class CorpusEntry:
    id: str
    source_docx: str
    expected_fields: dict
    known_fill_data: dict | None
    expected_output_lines: list[str]


@dataclass
class EntryResult:
    entry_id: str
    field_recall: float
    field_precision: float
    template_integrity: float
    round_trip_accuracy: float
    score: float
    errors: list[str]


# ── Corpus loading ───────────────────────────────────────────────────────

def load_corpus(subset: int | None = None) -> list[CorpusEntry]:
    """Load corpus entries from eval_corpus/corpus/{industry}/{template}/."""
    entries = []
    # Walk two levels: industry dirs contain template dirs
    entry_paths = []
    for industry in sorted(os.listdir(CORPUS_DIR)):
        industry_dir = os.path.join(CORPUS_DIR, industry)
        if not os.path.isdir(industry_dir):
            continue
        for template in sorted(os.listdir(industry_dir)):
            template_dir = os.path.join(industry_dir, template)
            if os.path.isdir(template_dir):
                entry_paths.append((f"{industry}/{template}", template_dir))

    if subset:
        entry_paths = entry_paths[:subset]

    for entry_id, entry_dir in entry_paths:
        source_docx = os.path.join(entry_dir, "source.docx")
        fields_path = os.path.join(entry_dir, "expected_fields.json")

        if not os.path.exists(source_docx) or not os.path.exists(fields_path):
            continue

        with open(fields_path) as f:
            expected_fields = json.load(f)

        fill_path = os.path.join(entry_dir, "known_fill_data.json")
        known_fill_data = None
        if os.path.exists(fill_path):
            with open(fill_path) as f:
                known_fill_data = json.load(f)

        output_path = os.path.join(entry_dir, "expected_output.txt")
        expected_output_lines = []
        if os.path.exists(output_path):
            with open(output_path) as f:
                expected_output_lines = [
                    line.strip() for line in f if line.strip()
                ]

        entries.append(CorpusEntry(
            id=entry_id,
            source_docx=source_docx,
            expected_fields=expected_fields,
            known_fill_data=known_fill_data,
            expected_output_lines=expected_output_lines,
        ))

    return entries


# ── Metric: Field Recall ─────────────────────────────────────────────────

def compute_field_recall(
    predicted_fields: list[dict],
    expected: dict,
) -> float:
    """
    Measure what fraction of expected fields were found.
    Returns adjusted recall in [0.0, 1.0].
    """
    expected_fields = expected["fields"]
    if not expected_fields:
        return 1.0

    # Build lookup: original_value -> expected field
    expected_by_value = {}
    expected_by_name = {}
    for ef in expected_fields:
        expected_by_value[ef["original_value"]] = ef
        expected_by_name[ef["field_name"]] = ef

    # Match predicted fields to expected fields
    matched_fields = set()
    matched_variations = 0
    total_variations = sum(len(ef.get("variations", [])) for ef in expected_fields)

    for pf in predicted_fields:
        pf_value = pf.get("original_value", "")
        pf_name = pf.get("field_name", "")

        # Try matching by original_value first, then by field_name
        ef = expected_by_value.get(pf_value) or expected_by_name.get(pf_name)
        if ef is None:
            continue

        matched_fields.add(ef["field_name"])

        # Check variation coverage
        predicted_var_texts = {
            v["text"] for v in pf.get("all_variations", [])
        }
        for ev in ef.get("variations", []):
            if ev["text"] in predicted_var_texts:
                matched_variations += 1

    field_recall = len(matched_fields) / len(expected_fields)
    variation_recall = (
        matched_variations / total_variations if total_variations > 0 else 1.0
    )

    R = 0.7 * field_recall + 0.3 * variation_recall

    # Critical field penalty
    critical_fields = [f for f in expected_fields if f.get("is_critical", False)]
    if critical_fields:
        missed_critical = sum(
            1 for f in critical_fields if f["field_name"] not in matched_fields
        )
        critical_miss_ratio = missed_critical / len(critical_fields)
        R = R * (1 - 0.3 * critical_miss_ratio)

    return R


# ── Metric: Field Precision ──────────────────────────────────────────────

def compute_field_precision(
    predicted_fields: list[dict],
    expected: dict,
) -> float:
    """
    Measure what fraction of predicted fields are actual variables (not boilerplate).
    Returns precision in [0.0, 1.0].
    """
    if not predicted_fields:
        return 1.0

    boilerplate = [
        bp.lower() for bp in expected.get("boilerplate_phrases", [])
    ]
    expected_values = {
        f["original_value"] for f in expected["fields"]
    }
    expected_names = {
        f["field_name"] for f in expected["fields"]
    }

    true_positives = 0
    false_positives = 0

    for pf in predicted_fields:
        pf_value = pf.get("original_value", "")
        pf_name = pf.get("field_name", "")

        # Check if this is a known expected field
        if pf_value in expected_values or pf_name in expected_names:
            true_positives += 1
            continue

        # Check if it looks like boilerplate
        is_boilerplate = any(bp in pf_value.lower() for bp in boilerplate)
        is_too_short = len(pf_value) < 3

        if is_boilerplate or is_too_short:
            false_positives += 1
        else:
            # Unknown field — not in expected, not obviously boilerplate.
            # Give partial credit: it might be a valid field we didn't annotate.
            true_positives += 0.5
            false_positives += 0.5

    total = true_positives + false_positives
    return true_positives / total if total > 0 else 1.0


# ── Metric: Template Integrity ───────────────────────────────────────────

def compute_template_integrity(template_path: str) -> float:
    """
    Check if the template loads and renders without errors.
    Returns score in [0.0, 1.0].
    """
    from docxtpl import DocxTemplate

    try:
        tpl = DocxTemplate(template_path)
    except Exception:
        return 0.0

    # Extract all placeholder names
    try:
        doc = Document(template_path)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += "\n" + cell.text
        placeholders = set(_PLACEHOLDER_RE.findall(all_text))
    except Exception:
        return 0.3

    # Build dummy context
    dummy_context = {name: "TEST_VALUE" for name in placeholders}

    try:
        tpl.render(dummy_context)
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=True) as tmp:
            tpl.save(tmp.name)
        return 1.0
    except Exception:
        return 0.3


# ── Metric: Round-Trip Accuracy ──────────────────────────────────────────

def compute_round_trip(
    artifacts: TemplateArtifacts,
    known_fill_data: dict,
    expected_lines: list[str],
    llm_config: LLMConfig,
) -> float:
    """
    Fill template with known data, check if expected phrases appear in output.
    Returns accuracy in [0.0, 1.0].
    """
    if not expected_lines:
        return 1.0

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        output_path = tmp.name

    try:
        result = fill_template(
            artifacts,
            new_facts="",
            output_path=output_path,
            llm_config=llm_config,
            fill_data_override=known_fill_data,
        )

        # Extract text from rendered document
        doc = Document(output_path)
        rendered_text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    rendered_text += "\n" + cell.text
        for section in doc.sections:
            for p in section.header.paragraphs + section.footer.paragraphs:
                rendered_text += "\n" + p.text

        # Check each expected line
        matched = sum(
            1 for line in expected_lines if line in rendered_text
        )
        return matched / len(expected_lines)

    except Exception:
        return 0.0
    finally:
        if os.path.exists(output_path):
            os.unlink(output_path)


# ── Full evaluation ──────────────────────────────────────────────────────

def evaluate_entry(
    entry: CorpusEntry,
    llm_config: LLMConfig,
) -> EntryResult:
    """Run full evaluation on one corpus entry. Returns per-entry result."""
    errors = []

    with tempfile.TemporaryDirectory() as tmp_dir:
        # Step 1: Run create_template (1 LLM call)
        try:
            artifacts = create_template(
                entry.source_docx, tmp_dir, llm_config
            )
        except Exception as e:
            errors.append(f"create_template failed: {e}")
            return EntryResult(
                entry_id=entry.id,
                field_recall=0.0,
                field_precision=0.0,
                template_integrity=0.0,
                round_trip_accuracy=0.0,
                score=0.0,
                errors=errors,
            )

        # Step 2: Field recall
        predicted_fields = artifacts.analysis.get("fields", [])
        recall = compute_field_recall(predicted_fields, entry.expected_fields)

        # Step 3: Field precision
        precision = compute_field_precision(
            predicted_fields, entry.expected_fields
        )

        # Step 4: Template integrity
        integrity = compute_template_integrity(artifacts.template_path)

        # Step 5: Round-trip accuracy
        if entry.known_fill_data and entry.expected_output_lines:
            roundtrip = compute_round_trip(
                artifacts,
                entry.known_fill_data,
                entry.expected_output_lines,
                llm_config,
            )
        else:
            roundtrip = 1.0

    # Composite score
    score = (
        WEIGHT_RECALL * recall
        + WEIGHT_PRECISION * precision
        + WEIGHT_INTEGRITY * integrity
        + WEIGHT_ROUNDTRIP * roundtrip
    )

    return EntryResult(
        entry_id=entry.id,
        field_recall=recall,
        field_precision=precision,
        template_integrity=integrity,
        round_trip_accuracy=roundtrip,
        score=score,
        errors=errors,
    )


# ── Fast check ───────────────────────────────────────────────────────────

def fast_check() -> bool:
    """
    Free, no-LLM checks. Returns True if all pass.
    Run before every commit / full eval.
    """
    print("=== Fast Check ===")
    passed = True

    # 1. Import check
    try:
        import docx_template_pipeline  # noqa: F401
        print("PASS  import docx_template_pipeline")
    except Exception as e:
        print(f"FAIL  import: {e}")
        passed = False

    # 2. Schema validation
    try:
        from docx_template_pipeline import _ANALYSIS_SCHEMA
        assert "properties" in _ANALYSIS_SCHEMA
        assert "fields" in _ANALYSIS_SCHEMA["properties"]
        print("PASS  _ANALYSIS_SCHEMA is valid")
    except Exception as e:
        print(f"FAIL  schema: {e}")
        passed = False

    # 3. Regex check
    try:
        from docx_template_pipeline import _PLACEHOLDER_RE
        assert _PLACEHOLDER_RE.search("{{ field_name }}")
        assert _PLACEHOLDER_RE.search("{{field_name}}")
        assert not _PLACEHOLDER_RE.search("plain text")
        print("PASS  _PLACEHOLDER_RE works correctly")
    except Exception as e:
        print(f"FAIL  regex: {e}")
        passed = False

    # 4. Template rendering check (cached templates only)
    corpus = load_corpus()
    for entry in corpus:
        cached_template = os.path.join(
            os.path.dirname(entry.source_docx), "cached_template.docx"
        )
        if os.path.exists(cached_template):
            integrity = compute_template_integrity(cached_template)
            status = "PASS" if integrity >= 0.7 else "FAIL"
            print(f"{status}  template integrity [{entry.id}]: {integrity:.2f}")
            if integrity < 0.7:
                passed = False

    print(f"\n{'ALL CHECKS PASSED' if passed else 'SOME CHECKS FAILED'}")
    return passed


# ── Main ─────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="DocForge evaluation — produces a scalar score (0.0-1.0)"
    )
    parser.add_argument(
        "--fast", action="store_true",
        help="Run fast checks only (no LLM calls, no score)"
    )
    parser.add_argument(
        "--subset", type=int, default=None,
        help="Evaluate only the first N corpus entries"
    )
    parser.add_argument(
        "--model", type=str, default=None,
        help="LLM model to use for extraction"
    )
    args = parser.parse_args()

    if args.fast:
        ok = fast_check()
        sys.exit(0 if ok else 1)

    # Load API config from .env
    from dotenv import load_dotenv
    env_path = os.path.join(os.path.dirname(__file__), "..", ".env")
    load_dotenv(env_path)

    api_key = os.environ.get("OPENROUTER_API_KEY") or os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        print("ERROR: Set OPENROUTER_API_KEY or ANTHROPIC_API_KEY", file=sys.stderr)
        sys.exit(1)

    model = args.model or "anthropic/claude-sonnet-4-20250514"
    base_url = "https://openrouter.ai/api" if os.environ.get("OPENROUTER_API_KEY") else None

    llm_config = LLMConfig(api_key=api_key, model=model, base_url=base_url)
    corpus = load_corpus(subset=args.subset)

    if not corpus:
        print("ERROR: No corpus entries found in eval_corpus/corpus/", file=sys.stderr)
        sys.exit(1)

    print(f"Evaluating {len(corpus)} corpus entries...\n")
    t0 = time.time()
    results = []
    error_count = 0

    for entry in corpus:
        print(f"  [{entry.id}] ", end="", flush=True)
        result = evaluate_entry(entry, llm_config)
        results.append(result)
        if result.errors:
            error_count += len(result.errors)
        status = "OK" if not result.errors else f"ERRORS: {result.errors}"
        print(f"score={result.score:.4f}  {status}")

    elapsed = time.time() - t0

    # Aggregate: mean across corpus entries
    avg_score = sum(r.score for r in results) / len(results)
    avg_recall = sum(r.field_recall for r in results) / len(results)
    avg_precision = sum(r.field_precision for r in results) / len(results)
    avg_integrity = sum(r.template_integrity for r in results) / len(results)
    avg_roundtrip = sum(r.round_trip_accuracy for r in results) / len(results)

    # Code bloat penalty — penalizes prompt-stuffing and unnecessary complexity
    pipeline_path = os.path.join(os.path.dirname(__file__), "..", "docx_template_pipeline.py")
    with open(pipeline_path) as f:
        current_lines = len(f.readlines())
    bloat_ratio = current_lines / BASELINE_LINES
    bloat_applied = False
    if bloat_ratio > BLOAT_THRESHOLD:
        avg_score *= BLOAT_PENALTY
        bloat_applied = True

    # Print summary in autoresearch-compatible format
    print("\n---")
    print(f"eval_score:         {avg_score:.6f}")
    print(f"field_precision:    {avg_precision:.4f}")
    print(f"field_recall:       {avg_recall:.4f}")
    print(f"template_accuracy:  {avg_integrity:.4f}")
    print(f"fill_accuracy:      {avg_roundtrip:.4f}")
    print(f"corpus_size:        {len(corpus)}")
    print(f"errors:             {error_count}")
    print(f"code_lines:         {current_lines} / {BASELINE_LINES} baseline ({bloat_ratio:.2f}x){' [BLOAT PENALTY APPLIED]' if bloat_applied else ''}")
    print(f"total_seconds:      {elapsed:.1f}")


if __name__ == "__main__":
    main()
