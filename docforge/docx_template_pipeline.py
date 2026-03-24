"""General-purpose document template pipeline.

Extracts variable fields from any .docx using an LLM, creates a reusable
template, and fills it with new data via LLM structured output.

Public API:
    create_template()    — analyze a .docx and build a reusable template + schema
    fill_template()      — fill the template with new data (via LLM or manual JSON)
    generate_fill_prompt() — get a copy-paste prompt for manual LLM use
"""

from __future__ import annotations

import json
import logging
import os
import re
import shutil
import warnings
from dataclasses import dataclass, field
from xml.sax.saxutils import escape as _xml_escape
from typing import Iterator

try:
    import anthropic
except ImportError:
    anthropic = None  # type: ignore[assignment]

from docx import Document
from docxtpl import DocxTemplate

__all__ = [
    "LLMConfig",
    "TemplateArtifacts",
    "RenderedDocument",
    "create_template",
    "fill_template",
    "generate_fill_prompt",
]

logger = logging.getLogger(__name__)

# ── Constants ────────────────────────────────────────────────────────────────

_PLACEHOLDER_RE = re.compile(r"\{\{\s*(\w+)\s*\}\}")

_TR_UPPER = "upper"
_TR_LOWER = "lower"
_TR_SHORT = "short"
_TR_ABBREV = "abbrev"
_TR_IDENTITY = "identity"
_TR_LASTNAME_AFFIX = "lastname_affix"
# Affix transform is a tuple: ("affix", prefix, suffix)

_USER_SUPPLIED_SUFFIXES = {"_short", "_abbrev"}

_ANALYSIS_SCHEMA: dict = {
    "type": "object",
    "required": ["document_type", "document_description", "fields"],
    "additionalProperties": False,
    "properties": {
        "document_type": {
            "type": "string",
            "description": "What type of document this is, e.g. 'Legal Motion', 'Contract', 'Business Letter'",
        },
        "document_description": {
            "type": "string",
            "description": "One-paragraph description of the document's purpose and structure",
        },
        "fields": {
            "type": "array",
            "description": "All variable/case-specific fields found in the document",
            "items": {
                "type": "object",
                "required": [
                    "field_name",
                    "original_value",
                    "category",
                    "description",
                    "all_variations",
                ],
                "additionalProperties": False,
                "properties": {
                    "field_name": {
                        "type": "string",
                        "description": "snake_case variable name, e.g. 'plaintiff_name', 'hearing_date'",
                    },
                    "original_value": {
                        "type": "string",
                        "description": "The canonical/primary text value to replace. Must match the document EXACTLY.",
                    },
                    "all_variations": {
                        "type": "array",
                        "description": "ALL variations of this value (different casing, abbreviations, prefixed forms). Each must be EXACT verbatim. Order longest to shortest.",
                        "items": {
                            "type": "object",
                            "required": ["text", "placeholder"],
                            "additionalProperties": False,
                            "properties": {
                                "text": {
                                    "type": "string",
                                    "description": "Exact text as it appears in the document",
                                },
                                "placeholder": {
                                    "type": "string",
                                    "description": "Jinja2 placeholder, e.g. '{{ plaintiff_name }}' or '{{ plaintiff_name_upper }}'",
                                },
                            },
                        },
                    },
                    "category": {
                        "type": "string",
                        "description": "Field category for grouping, e.g. 'attorney', 'parties', 'dates', 'addresses'",
                    },
                    "description": {
                        "type": "string",
                        "description": "Human-readable description of what this field represents.",
                    },
                },
            },
        },
    },
}

_ANALYSIS_SYSTEM_PROMPT = """\
You are a document template analyst. Given a document's full text, you identify \
every field that is specific to this particular instance and would change if the \
document were reused for a different case/client/situation.

CRITICAL RULES:
1. The 'original_value' and each variation 'text' must be EXACT verbatim matches from the document.
   Copy-paste precision. If you get the text wrong, the template will break.
2. Identify ALL variations of the same entity (uppercase, title case, abbreviated, with/without
   suffixes like LLC, etc.). Each gets its own placeholder suffix like _upper, _short, _abbrev.
3. Order variations from LONGEST to SHORTEST to prevent partial-match corruption.
4. DO NOT mark boilerplate language, statute citations, or standard phrases as variable fields.
   Only mark content specific to THIS case/client.
5. Be thorough — missing a field means it stays hardcoded in the template.
6. Group logically: attorney info, party names, case details, court info, dates, addresses, people, etc.
7. For EVERY person identified, you MUST also capture ALL reference forms used in the body:
   - Standalone last name (e.g., "Smith" when full name is "John Smith") → suffix _lastname
   - Compound references like "LastName Decl.", "LastName Declaration" → suffix _lastname_decl
   - These are extremely common in legal/business docs. Scan the ENTIRE body for every surname.
8. For party names, ALWAYS capture the short abbreviation/acronym if used in the body
   (e.g., "ASP", "GSO"). Use suffix _short. If missed, abbreviations stay hardcoded.
9. When a value appears WITH a prefix like "Hon." or "SBN" in some places but WITHOUT it
   in others, capture BOTH forms as separate variations. Use suffix _with_hon, _with_prefix, etc.
10. NEVER create a field variation with a value shorter than 3 characters (e.g., "2", "CA").
    Short values cause false-positive replacements in section numbers, bullet points, and other
    unrelated text. Instead, include short values as part of a longer surrounding phrase."""


# ── Dataclasses ──────────────────────────────────────────────────────────────


@dataclass
class LLMConfig:
    """Configuration for the LLM client."""

    api_key: str
    model: str = "anthropic/claude-3.5-sonnet"
    base_url: str | None = None
    max_tokens_analyze: int = 8000
    max_tokens_fill: int = 4000


@dataclass
class TemplateArtifacts:
    """Outputs from create_template(). Pass to fill_template() to generate documents."""

    analysis: dict
    fill_schema: dict
    template_path: str
    document_type: str
    document_description: str
    field_count: int
    placeholder_count: int


@dataclass
class RenderedDocument:
    """Outputs from fill_template()."""

    output_path: str
    fill_data: dict
    context: dict
    placeholder_count: int


# ── Private utilities ────────────────────────────────────────────────────────


def _make_client(config: LLMConfig):
    if anthropic is None:
        raise ImportError("anthropic package is required: pip install anthropic")
    kwargs: dict = {"api_key": config.api_key}
    if config.base_url:
        kwargs["base_url"] = config.base_url
    return anthropic.Anthropic(**kwargs)


def _extract_placeholder_name(placeholder_str: str) -> str | None:
    m = _PLACEHOLDER_RE.search(placeholder_str)
    return m.group(1) if m else None


def _iter_field_variations(analysis: dict) -> Iterator[tuple[dict, dict, str]]:
    for fld in analysis["fields"]:
        for v in fld["all_variations"]:
            pname = _extract_placeholder_name(v["placeholder"])
            if pname:
                yield fld, v, pname


def _apply_cross_replacements(text: str, cross_map: list[tuple[str, str]]) -> str:
    for orig_val, new_val in cross_map:
        text = text.replace(orig_val, new_val)
    return text


# ── Phase 1: Extract ─────────────────────────────────────────────────────────


def _extract_document_text(docx_path: str) -> str:
    doc = Document(docx_path)
    parts: list[str] = []

    for section in doc.sections:
        for p in section.header.paragraphs:
            if p.text.strip():
                parts.append(f"[HEADER] {p.text.strip()}")
        for p in section.footer.paragraphs:
            if p.text.strip():
                parts.append(f"[FOOTER] {p.text.strip()}")

    for i, table in enumerate(doc.tables):
        parts.append(f"\n[TABLE {i}]")
        for r, row in enumerate(table.rows):
            for c, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text:
                    text = re.sub(r"\n+", " | ", text)
                    parts.append(f"  Row {r}, Col {c}: {text}")

    parts.append("\n[BODY]")
    for i, p in enumerate(doc.paragraphs):
        if not p.text.strip():
            continue
        style = p.style.name if p.style else "Normal"
        parts.append(f"  P{i} [{style}]: {p.text.strip()}")

    return "\n".join(parts)


# ── Phase 2: Analyze ─────────────────────────────────────────────────────────


def _analyze_document(doc_text: str, config: LLMConfig) -> dict:
    client = _make_client(config)
    response = client.messages.create(
        model=config.model,
        max_tokens=config.max_tokens_analyze,
        system=_ANALYSIS_SYSTEM_PROMPT,
        tools=[
            {
                "name": "report_fields",
                "description": "Report all variable fields found in the document",
                "input_schema": _ANALYSIS_SCHEMA,
            }
        ],
        tool_choice={"type": "tool", "name": "report_fields"},
        messages=[
            {
                "role": "user",
                "content": f"""Analyze this document and identify every variable/case-specific field.

DOCUMENT TEXT:
---
{doc_text}
---

Find ALL fields that would change for a different case. Be exhaustive.

IMPORTANT REMINDERS:
- For each person, search the ENTIRE body for standalone last-name references and compound
  forms (e.g., "Manzanares Decl.", just "Manzanares"). These MUST appear in all_variations.
- Party abbreviations (e.g., "ASP", "GSO") used throughout the body MUST be captured as _short variations.
- Prefixed forms like "Hon. Judge Name" or "SBN 12345" MUST be captured alongside the bare value.""",
            }
        ],
    )

    for block in response.content:
        if block.type == "tool_use":
            return block.input
    raise ValueError("LLM response did not contain a tool_use block")


# ── Phase 3: Build template ──────────────────────────────────────────────────


_MIN_REPLACEMENT_LENGTH = 3  # Skip values shorter than this to avoid false positives


def _build_replacement_rules(analysis: dict) -> list[tuple[str, str]]:
    rules = []
    for _, v, _ in _iter_field_variations(analysis):
        text = v["text"]
        if len(text) < _MIN_REPLACEMENT_LENGTH:
            logger.warning(
                "Skipping short replacement value '%s' (len=%d) — too likely to cause "
                "false-positive replacements in section numbers, bullet points, etc.",
                text,
                len(text),
            )
            continue
        rules.append((text, v["placeholder"]))
    rules.sort(key=lambda x: len(x[0]), reverse=True)
    return rules


def _apply_replacements_to_docx(
    source_path: str, output_path: str, rules: list[tuple[str, str]]
) -> None:
    shutil.copy2(source_path, output_path)
    doc = Document(output_path)

    def replace_in_paragraph(paragraph):
        if not paragraph.runs:
            return
        full_text = "".join(r.text for r in paragraph.runs)
        original = full_text
        for old, new in rules:
            full_text = full_text.replace(old, new)
        if full_text != original:
            paragraph.runs[0].text = full_text
            for run in paragraph.runs[1:]:
                run.text = ""

    for para in doc.paragraphs:
        replace_in_paragraph(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para)
    for section in doc.sections:
        for para in section.header.paragraphs:
            replace_in_paragraph(para)
        for para in section.footer.paragraphs:
            replace_in_paragraph(para)

    doc.save(output_path)


# ── Phase 4: Build fill schema ───────────────────────────────────────────────


def _build_fill_schema(analysis: dict) -> dict:
    categories: dict[str, list] = {}
    for fld in analysis["fields"]:
        categories.setdefault(fld["category"], []).append(fld)

    properties = {}
    required = []

    for cat, fields in categories.items():
        cat_properties: dict = {}
        cat_required: list = []

        for fld in fields:
            fname = fld["field_name"]
            cat_properties[fname] = {"type": "string", "description": fld["description"]}
            cat_required.append(fname)

            for _, v, pname in _iter_field_variations({"fields": [fld]}):
                for suffix in _USER_SUPPLIED_SUFFIXES:
                    if pname.endswith(suffix) and pname not in cat_properties:
                        cat_properties[pname] = {
                            "type": "string",
                            "description": (
                                f"Short form/abbreviation for {fld['description'].lower()} "
                                f"(in the original document this was '{v['text']}')"
                            ),
                        }
                        cat_required.append(pname)

        properties[cat] = {
            "type": "object",
            "required": cat_required,
            "additionalProperties": False,
            "properties": cat_properties,
        }
        required.append(cat)

    return {
        "type": "object",
        "required": required,
        "additionalProperties": False,
        "properties": properties,
    }


# ── Phase 5: Fill + Render ───────────────────────────────────────────────────


def _build_variation_map(analysis: dict) -> dict:
    vmap: dict[str, list] = {}
    for fld, v, pname in _iter_field_variations(analysis):
        fname = fld["field_name"]
        original_value = fld["original_value"]
        variation_text = v["text"]

        if fname not in vmap:
            vmap[fname] = []

        suffix_part = pname[len(fname) :] if pname.startswith(fname) else ""

        if pname.endswith("_upper"):
            transform = _TR_UPPER
        elif pname.endswith("_lower"):
            transform = _TR_LOWER
        elif pname.endswith("_short"):
            transform = _TR_SHORT
        elif pname.endswith("_abbrev"):
            transform = _TR_ABBREV
        elif "_lastname" in suffix_part:
            transform = _TR_LASTNAME_AFFIX
        elif pname == fname:
            transform = _TR_IDENTITY
        else:
            idx = variation_text.find(original_value)
            if idx < 0:
                idx = variation_text.lower().find(original_value.lower())
            if idx >= 0:
                prefix = variation_text[:idx]
                suffix = variation_text[idx + len(original_value) :]
                transform = ("affix", prefix, suffix)
            else:
                transform = _TR_IDENTITY

        vmap[fname].append((pname, transform, variation_text, original_value))
    return vmap


def _build_cross_field_map(
    analysis: dict, fill_data: dict
) -> list[tuple[str, str]]:
    pairs = []
    for fld in analysis["fields"]:
        fname = fld["field_name"]
        cat = fld["category"]
        orig = fld["original_value"]
        new_val = fill_data.get(cat, {}).get(fname)
        if new_val and orig != new_val and len(orig) >= 3:
            pairs.append((orig, new_val))
    pairs.sort(key=lambda x: len(x[0]), reverse=True)
    return pairs


def _expand_to_placeholders(fill_data: dict, analysis: dict) -> dict:
    vmap = _build_variation_map(analysis)
    cross_map = _build_cross_field_map(analysis, fill_data)
    context: dict[str, str] = {}

    for fld in analysis["fields"]:
        fname = fld["field_name"]
        cat = fld["category"]
        canonical = fill_data.get(cat, {}).get(fname, "[TBD]")

        for pname, transform, _original, field_orig_value in vmap.get(fname, []):
            if transform == _TR_UPPER:
                context[pname] = canonical.upper()
            elif transform == _TR_LOWER:
                context[pname] = canonical.lower()
            elif transform == _TR_IDENTITY:
                context[pname] = canonical
            elif isinstance(transform, tuple) and transform[0] == "affix":
                _, prefix, suffix = transform
                context[pname] = _apply_cross_replacements(
                    f"{prefix}{canonical}{suffix}", cross_map
                )
            elif transform == _TR_LASTNAME_AFFIX:
                new_lastname = canonical.split()[-1] if canonical.strip() else canonical
                orig_lastname = (
                    field_orig_value.split()[-1]
                    if field_orig_value.strip()
                    else field_orig_value
                )
                idx = _original.find(orig_lastname)
                if idx >= 0:
                    prefix = _original[:idx]
                    suffix = _original[idx + len(orig_lastname) :]
                    context[pname] = _apply_cross_replacements(
                        f"{prefix}{new_lastname}{suffix}", cross_map
                    )
                else:
                    context[pname] = new_lastname
            elif transform in (_TR_SHORT, _TR_ABBREV):
                alt_val = fill_data.get(cat, {}).get(f"{fname}_{transform}")
                if not alt_val:
                    alt_val = fill_data.get(cat, {}).get(pname)
                context[pname] = alt_val if alt_val else canonical
            else:
                warnings.warn(
                    f"Unknown transform '{transform}' for placeholder '{pname}', "
                    "using canonical value"
                )
                context[pname] = canonical

    return context


def _fill_with_llm(
    new_facts: str, analysis: dict, fill_schema: dict, config: LLMConfig
) -> dict:
    doc_type = analysis["document_type"]
    doc_desc = analysis["document_description"]

    schema_fields = []
    for cat, spec in fill_schema["properties"].items():
        for fname, fspec in spec["properties"].items():
            schema_fields.append(f"- {fname} ({cat}): {fspec['description']}")

    client = _make_client(config)
    response = client.messages.create(
        model=config.model,
        max_tokens=config.max_tokens_fill,
        system=f"""You are a document data assistant. You produce structured JSON to fill
a {doc_type} template.

Document description: {doc_desc}

Fields to fill (from schema):
{chr(10).join(schema_fields)}

Rules:
- Use ONLY the facts provided. Never fabricate information.
- If a fact is not provided, use a reasonable placeholder like '[TBD]'.
- Dates should be in 'Month Day, Year' format (e.g., 'January 16, 2026').
- Provide the canonical/primary form of each value. Variations (uppercase, lowercase)
  will be derived automatically.
- For _short fields, provide the abbreviation/acronym (e.g., 'PSC', 'GGL').""",
        tools=[
            {
                "name": "fill_template",
                "description": f"Provide structured data to fill a {doc_type} template",
                "input_schema": fill_schema,
            }
        ],
        tool_choice={"type": "tool", "name": "fill_template"},
        messages=[
            {
                "role": "user",
                "content": f"Fill in the template fields based on these facts:\n\n{new_facts}",
            }
        ],
    )

    for block in response.content:
        if block.type == "tool_use":
            return block.input
    raise ValueError("LLM response did not contain a tool_use block")


# ── Public API ───────────────────────────────────────────────────────────────


def create_template(
    source_docx: str,
    output_dir: str,
    llm_config: LLMConfig,
    *,
    template_filename: str = "template.docx",
) -> TemplateArtifacts:
    """Analyze a .docx and build a reusable template + fill schema.

    Args:
        source_docx: Path to the source .docx file.
        output_dir: Directory where the template .docx will be written.
        llm_config: LLM connection configuration.
        template_filename: Name for the generated template file.

    Returns:
        TemplateArtifacts containing the analysis, schema, and template path.
    """
    source_docx = os.path.abspath(source_docx)
    output_dir = os.path.abspath(output_dir)

    if not os.path.isfile(source_docx):
        raise FileNotFoundError(f"Source document not found: {source_docx}")

    os.makedirs(output_dir, exist_ok=True)
    template_path = os.path.join(output_dir, template_filename)

    logger.info("Phase 1: Extracting document text")
    doc_text = _extract_document_text(source_docx)

    logger.info("Phase 2: Analyzing document with LLM")
    analysis = _analyze_document(doc_text, llm_config)

    logger.info("Phase 3: Building template")
    rules = _build_replacement_rules(analysis)
    _apply_replacements_to_docx(source_docx, template_path, rules)

    logger.info("Phase 4: Building fill schema")
    fill_schema = _build_fill_schema(analysis)

    # Count placeholders in the template
    doc = Document(template_path)
    all_parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_parts.append(cell.text)
    for section in doc.sections:
        for p in section.header.paragraphs + section.footer.paragraphs:
            all_parts.append(p.text)
    placeholder_count = len(set(_PLACEHOLDER_RE.findall("\n".join(all_parts))))

    return TemplateArtifacts(
        analysis=analysis,
        fill_schema=fill_schema,
        template_path=template_path,
        document_type=analysis["document_type"],
        document_description=analysis["document_description"],
        field_count=len(analysis["fields"]),
        placeholder_count=placeholder_count,
    )


def fill_template(
    artifacts: TemplateArtifacts,
    new_facts: str,
    output_path: str,
    llm_config: LLMConfig,
    *,
    fill_data_override: dict | None = None,
) -> RenderedDocument:
    """Fill the template with new data and render a .docx.

    Args:
        artifacts: Output from create_template().
        new_facts: Plain-text description of the new case/data.
        output_path: Path for the generated .docx file.
        llm_config: LLM connection configuration (unused if fill_data_override is provided).
        fill_data_override: Skip the LLM call and use this data directly.

    Returns:
        RenderedDocument with the output path and fill data used.
    """
    output_path = os.path.abspath(output_path)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    if fill_data_override is not None:
        fill_data = fill_data_override
        logger.info("Using provided fill data (skipping LLM call)")
    else:
        logger.info("Generating fill data with LLM")
        fill_data = _fill_with_llm(
            new_facts, artifacts.analysis, artifacts.fill_schema, llm_config
        )

    context = _expand_to_placeholders(fill_data, artifacts.analysis)

    # XML-escape special characters (&, <, >) so docxtpl renders them correctly
    context = {k: _xml_escape(v) if isinstance(v, str) else v for k, v in context.items()}

    tpl = DocxTemplate(artifacts.template_path)
    tpl.render(context)
    tpl.save(output_path)

    logger.info("Document generated: %s (%d placeholders filled)", output_path, len(context))

    return RenderedDocument(
        output_path=output_path,
        fill_data=fill_data,
        context=context,
        placeholder_count=len(context),
    )


def generate_fill_prompt(artifacts: TemplateArtifacts, new_facts: str) -> str:
    """Generate a copy-paste prompt for manual LLM use.

    Paste the returned string into any LLM chat (Claude, ChatGPT, etc.),
    copy the JSON response, and pass it as fill_data_override to fill_template().
    """
    field_list = "\n".join(
        f"- {f['field_name']} ({f['category']}): {f['description']}"
        for f in artifacts.analysis["fields"]
    )

    return f"""You produce structured JSON to fill a {artifacts.document_type} template.

## Fields to fill:
{field_list}

## JSON Schema (conform EXACTLY):
```json
{json.dumps(artifacts.fill_schema, indent=2)}
```

## Rules:
- Use ONLY the facts below. Never fabricate.
- Dates: 'Month Day, Year' format.
- If a fact is missing, use '[TBD]'.

## Facts:
{new_facts}

Return ONLY the JSON. No markdown fences. No explanation."""
